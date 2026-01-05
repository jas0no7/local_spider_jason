import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER


# =========================
# 1. 整页背景图（页眉方式）
# =========================
def add_full_page_background(document, image_path):
    section = document.sections[0]
    header = section.header
    header.is_linked_to_previous = False

    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run()
    run.add_picture(
        image_path,
        width=section.page_width,
        height=section.page_height
    )


def clear_header(document):
    section = document.sections[0]
    header = section.header
    for p in header.paragraphs:
        p.clear()


# =========================
# 2. 构造期号行文本
# =========================
def build_issue_line(year, issue_no, start, end, total_no):
    left = f"{year}年 第{issue_no}期（{start.strftime('%m.%d')}-{end.strftime('%m.%d')}）"
    right = f"总第{total_no}期"
    return left, right


def add_issue_line(document, year, issue_no, start, end, total_no):
    left, right = build_issue_line(year, issue_no, start, end, total_no)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # 右对齐制表位（控制“总第XX期”的位置）
    tabs = p.paragraph_format.tab_stops
    tabs.add_tab_stop(
        Pt(450),  # 可微调：420~480
        WD_TAB_ALIGNMENT.RIGHT,
        WD_TAB_LEADER.SPACES
    )

    run = p.add_run(f"{left}\t{right}")
    run.font.size = Pt(12)
    run.font.bold = False


# =========================
# 3. 封面函数
# =========================
def add_cover(
    document,
    bg_image_path,
    title_main,
    title_sub,
    year,
    issue_no,
    start,
    end,
    total_no
):
    # 背景图
    add_full_page_background(document, bg_image_path)

    # 上部留白（配合背景图调）
    for _ in range(4):
        document.add_paragraph()

    # 主标题（职教周刊）
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(title_main)
    run.font.size = Pt(26)
    run.font.bold = True

    # 副标题（高职）
    if title_sub:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(title_sub)
        run.font.size = Pt(14)
        run.font.bold = False

    document.add_paragraph()  # 空一行

    # 动态期号行（重点）
    add_issue_line(document, year, issue_no, start, end, total_no)

    # 下部留白
    for _ in range(12):
        document.add_paragraph()

    # 发文单位（可按需要删/改）
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("高职发展智库")
    run.font.size = Pt(12)
    run.font.bold = False

    # 封面结束
    document.add_page_break()
    clear_header(document)


# =========================
# 4. 主程序（本地测试）
# =========================
def main():
    doc = Document()

    # ====== 你要改的参数 ======
    bg_image_path = r"D:\项目文件夹\local_spider_jason\职教大脑\cover_bg.png"

    title_main = "职教周刊"
    title_sub = "（高职）"

    year = 2026
    issue_no = 1
    total_no = 50

    start = datetime.date(2025, 12, 29)
    end = datetime.date(2026, 1, 4)
    # =========================

    add_cover(
        document=doc,
        bg_image_path=bg_image_path,
        title_main=title_main,
        title_sub=title_sub,
        year=year,
        issue_no=issue_no,
        start=start,
        end=end,
        total_no=total_no
    )

    output = "职教周刊_封面测试.docx"
    doc.save(output)
    print(f"生成成功：{output}")


if __name__ == "__main__":
    main()
