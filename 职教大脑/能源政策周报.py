import os

from PIL import Image, ImageDraw, ImageFont
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt
from yd_llm_processing import YdLlmProcessing
from docx import Document, enum, shared
from docx.oxml.ns import qn
from docx.shared import Cm
# from report_docx_utils import write_title, write_docx
from report_llm_deal_content import get_ai_summary

from cron_scripts import common_es, common_sql
from common.tools import PublicTools

# 初始化文档
def document_init():
    document = Document()
    # 正文样式
    document.styles['Normal'].font.name = 'FZFangSong-Z02'  # 西文字符
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '方正仿宋_GBK')  # 中文字符
    document.styles['Normal'].paragraph_format.space_before = shared.Pt(0)  # 段落设置-段前
    document.styles['Normal'].paragraph_format.space_after = shared.Pt(0)  # 段落设置-段后
    # 获取第一节
    section = document.sections[0]

    section.footer.is_linked_to_previous = False
    if not section.footer.paragraphs:
        footer_para = section.footer.add_paragraph()
    else:
        footer_para = section.footer.paragraphs[0]

    footer_para.clear()
    footer_para.alignment = 1  # 居中

    # 添加页码并设置字体
    run = footer_para.add_run()

    # 创建页码字段
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._element.append(fldChar)

    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE'
    run._element.append(instrText)

    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    run._element.append(fldChar)

    # 设置字体
    run.font.name = '方正仿宋_GBK'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '方正仿宋_GBK')
    run.font.size = Pt(14)  # 四号字

    return document


# 文档添加封面
def add_cover(doc, report_date, week):
    def measure_text(draw1, text, font):
        """
        临时方法来测量文本尺寸
        """
        # 绘制文本（实际上并不会显示），只用来获取尺寸
        bbox = draw1.textbbox((0, 0), text, font=font)
        return bbox[2] - bbox[0], bbox[3] - bbox[1]
    text_color = (0, 121, 138)

    current_dir = os.path.dirname(os.path.abspath(__file__))
    img_path = os.path.join(current_dir, 'cover_week.png')
    img = Image.open(img_path)
    draw = ImageDraw.Draw(img)

    text1 = str(report_date.year)
    text2 = str(week)
    font_path = os.path.join(current_dir, 'arial.ttf')
    font_size = 100
    # 加载字体
    try:
        font = ImageFont.truetype(font_path, font_size)
    except IOError:
        print("字体文件未找到，请检查路径是否正确。")
        return

    text_width, text_height = measure_text(draw, text1, font)
    img_width, img_height = img.size
    x1 = (img_width - text_width) / 5.25
    y1 = (img_height - text_height) / 2.275

    draw.text((x1, y1), text1, fill=text_color, font=font)

    text_width, text_height = measure_text(draw, text2, font)
    img_width, img_height = img.size
    x2 = (img_width - text_width) / 3
    y2 = (img_height - text_height) / 2.275

    draw.text((x2, y2), text2, fill=text_color, font=font)

    new_image_p
