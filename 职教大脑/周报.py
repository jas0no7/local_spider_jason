import datetime
import calendar
from docx import Document, enum, shared
from docx.oxml.ns import qn


# =========================
# 1. Word 文档初始化
# =========================
def document_init():
    document = Document()
    style = document.styles['Normal']
    style.font.name = 'Times New Roman'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_before = shared.Pt(0)
    style.paragraph_format.space_after = shared.Pt(0)
    return document


# =========================
# 2. 封面
# =========================
def add_cover(document, title_str, start, stop, number):
    # 标题
    for i in range(6):
        p = document.add_paragraph(style='Normal')
        p.alignment = enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.add_run(title_str if i == 5 else "")
        run.font.size = shared.Pt(36)
        run.font.bold = True
        run.font.name = 'Times New Roman'   # 不再碰 rFonts

    # 期号
    phase = document.add_paragraph(style='Normal')
    phase.alignment = enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
    run = phase.add_run(
        f"{start.year}年第{number}期\n"
        f"{start.strftime('%m月%d日')} - {stop.strftime('%m月%d日')}"
    )
    run.font.size = shared.Pt(16)
    run.font.bold = True
    run.font.name = 'Times New Roman'

    # 空行
    for _ in range(10):
        document.add_paragraph()

    # 发文单位
    org = document.add_paragraph(style='Normal')
    org.alignment = enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
    run = org.add_run("广东省教育厅（职终处）\n深圳职业技术大学")
    run.font.size = shared.Pt(18)
    run.font.bold = True
    run.font.name = 'Times New Roman'

    document.add_page_break()



# =========================
# 3. 标题
# =========================
def write_title(document, level, content):
    title = document.add_heading('', level)
    title.alignment = enum.text.WD_PARAGRAPH_ALIGNMENT.LEFT

    run = title.add_run(content)

    # 只用官方 API，绝不碰 XML
    run.font.name = 'Times New Roman'
    run.font.bold = True

    if level == 1:
        run.font.size = shared.Pt(18)   # 小二
    elif level == 2:
        run.font.size = shared.Pt(15)   # 小三
    elif level == 3:
        run.font.size = shared.Pt(14)   # 四号
    else:
        run.font.size = shared.Pt(12)   # 小四



# =========================
# 4. 正文
# =========================
def write_paragraph(document, content, is_url=False):
    p = document.add_paragraph()
    p.alignment = enum.text.WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p.paragraph_format.first_line_indent = shared.Pt(24)
    run = p.add_run(content)
    run.font.size = shared.Pt(12)
    if is_url:
        run.font.color.rgb = shared.RGBColor(0, 0, 255)
        run.underline = True


# =========================
# 5. 模拟政策数据（死变量）
# =========================
MOCK_POLICIES = [
    {
        "title": "关于深化现代职业教育体系建设的意见",
        "publish_time": "2025-01-05",
        "ai_summary": "文件提出以产教融合为核心，完善现代职业教育体系。",
        "url": "https://example.com/policy1"
    },
    {
        "title": "推进职业院校数字化转型实施方案",
        "publish_time": "2025-01-10",
        "ai_summary": "方案强调以数字技术赋能教学管理和人才培养。",
        "url": "https://example.com/policy2"
    }
]


# =========================
# 6. 写政策模块
# =========================
def write_policy_section(document, title, policies):
    write_title(document, 1, title)
    for item in policies:
        date = datetime.datetime.strptime(item["publish_time"], "%Y-%m-%d").strftime("%m月%d日")
        write_title(document, 3, f"【{date}】{item['title']}")
        write_paragraph(document, item["ai_summary"])
        write_paragraph(document, f"原文链接：{item['url']}", is_url=True)
    document.add_page_break()


# =========================
# 7. 主流程（本地执行）
# =========================
def main():
    today = datetime.date.today()
    mid = datetime.date(today.year, today.month, 15)

    if today <= mid:
        start = datetime.date(today.year, today.month, 1)
        stop = mid
    else:
        start = datetime.date(today.year, today.month, 16)
        stop = datetime.date(
            today.year,
            today.month,
            calendar.monthrange(today.year, today.month)[1]
        )

    document = document_init()
    add_cover(document, "广东职业教育动态", start, stop, number=1)

    write_policy_section(document, "广东动态", MOCK_POLICIES)
    write_policy_section(document, "国家动态", MOCK_POLICIES)
    write_policy_section(document, "外省动态", MOCK_POLICIES)
    write_policy_section(document, "产教融合", MOCK_POLICIES)
    write_policy_section(document, "重大时评", MOCK_POLICIES)

    filename = "职业教育动态-本地测试.docx"
    document.save(filename)
    print(f"生成成功：{filename}")


if __name__ == "__main__":
    main()
