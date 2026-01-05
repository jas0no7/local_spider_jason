###
import os
import time
from typing import Dict, List

from docxtpl import DocxTemplate
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT


def _ensure_http(url: str) -> str:
    if not url:
        return ""
    url = url.strip()
    if not url:
        return ""
    if url.startswith(("http://", "https://")):
        return url
    return "https://" + url


def _make_hyperlink_run(paragraph, text: str, url: str):
    """
    在指定 paragraph 中创建一个超链接 run，显示 text，跳转 url。
    返回创建的 hyperlink 元素。
    """
    url = _ensure_http(url)
    if not url:
        return None

    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")

    r_style = OxmlElement("w:rStyle")
    r_style.set(qn("w:val"), "Hyperlink")
    r_pr.append(r_style)

    new_run.append(r_pr)

    w_t = OxmlElement("w:t")
    w_t.text = text
    new_run.append(w_t)

    hyperlink.append(new_run)
    return hyperlink


def _replace_text_with_hyperlink(paragraph, target_text: str, url: str) -> bool:
    """
    在一个段落中，把首次出现的 target_text 替换为超链接（显示 target_text，跳转 url）。
    返回是否替换成功。
    限制：需要 target_text 位于单个 run 内（这是最稳的情况）。
    """
    if not target_text:
        return False

    for run in paragraph.runs:
        if target_text in run.text:
            parts = run.text.split(target_text, 1)
            before = parts[0]
            after = parts[1]

            run.text = before

            hyperlink = _make_hyperlink_run(paragraph, target_text, url)
            if hyperlink is None:
                run.text = before + target_text + after
                return False

            run._r.addnext(hyperlink)

            if after:
                tail_run = paragraph.add_run(after)
                tail_run._r.addprevious(run._r.getnext())

            return True

    return False


def postprocess_add_source_hyperlinks(docx_path: str, items_by_source: List[Dict[str, str]]):
    """
    对生成好的 docx 做后处理：把（spider_from）里的 spider_from 文本变成超链接，跳转到 url。
    items_by_source: 每个元素包含 spider_from 和 url。
    """
    doc = Document(docx_path)

    source_to_url: Dict[str, str] = {}
    for it in items_by_source:
        spider_from = (it.get("spider_from") or "").strip()
        url = _ensure_http(it.get("url") or "")
        if spider_from and url and spider_from not in source_to_url:
            source_to_url[spider_from] = url

    if not source_to_url:
        doc.save(docx_path)
        return

    for p in doc.paragraphs:
        p_text = p.text
        if not p_text:
            continue

        for spider_from, url in source_to_url.items():
            if spider_from in p_text:
                _replace_text_with_hyperlink(p, spider_from, url)

    doc.save(docx_path)


def generate_weekly_report():
    template_path = "template.docx"
    output_path = "2026年第1期_生成测试.docx"

    if not os.path.exists(template_path):
        print(f"[错误] 找不到模板文件: {template_path}")
        return

    tpl = DocxTemplate(template_path)

    context = {
        "year": "2025",
        "issue": "1",
        "date_range": "12.29-01.04",
        "total_issue": "50",

        "section_zhiku": [
            {
                "title": "2025年高职院校发明专利授权量TOP100公布，江苏遥遥领先",
                "content": "近日，高职发展智库发布了2025年全国高职院校发明专利授权量统计数据。结果显示，江苏、浙江、广东高校表现优异，其中江苏省以总授权量3000件位居榜首。分析指出，高职院校科研成果转化率正逐年提升，服务区域经济能力显著增强。",
                "spider_from": "高职发展智库",
                "url": "https://www.example.com/"
            },
            {
                "title": "各省“双高计划”中期绩效评价结果出炉，优良率超85%",
                "content": "教育部及财政部联合发布通知，公布了双高计划的中期评价等级。评价结果显示，全国197所双高建设单位中，优秀等级占比为30%，良好等级占比为55%。部分院校在产教融合实训基地建设方面取得了突破性进展。",
                "spider_from": "教育部官网",
                "url": "https://www.example.com/"
            }
        ],

        "section_national": [
            {
                "title": "教育部：将急需紧缺专业纳入职业教育目录",
                "content": "为适应数字经济发展，教育部职成司表示将加快更新专业目录，增设人工智能、大数据分析等相关专业。此举旨在解决目前职教人才培养与产业需求脱节的问题，确保学生毕业即就业。",
                "spider_from": "中国教育报",
                "url": "https://www.example.com/"
            },
            {
                "title": "国家发改委：职业教育产教融合赋能提升行动实施方案",
                "content": "国家发改委发布最新实施方案，提出到2025年，国家产教融合试点城市达到50个左右，试点企业达到10000家以上。方案强调要给予产教融合型企业“金融+财政+土地+信用”的组合式激励。",
                "spider_from": "国家发改委",
                "url": "https://www.example.com/"
            }
        ],

        "section_provincial": [
            {
                "title": "广东省发布职业教育高质量发展三年行动计划",
                "content": "广东省厅发布最新文件，明确提出到2027年，全省建成20所具有国际影响力的职业技术大学。计划还指出，要加大对粤东粤西粤北地区职业教育的投入力度，实现全省职教均衡发展。",
                "spider_from": "广东省教育厅",
                "url": "https://www.example.com/"
            },
            {
                "title": "江苏省启动新一轮现代学徒制试点工作",
                "content": "江苏省教育厅宣布启动第五批现代学徒制试点，重点聚焦高端装备制造、新能源汽车等战略性新兴产业。试点项目将获得省级财政专项经费支持。",
                "spider_from": "江苏教育发布",
                "url": "www.example.com"  # 故意不带协议，测试自动补 https://
            }
        ],

        "section_college": [
            {
                "title": "深圳职业技术大学成立未来技术学院",
                "content": "1月6日，深职大未来技术学院揭牌仪式举行。该学院旨在培养掌握未来关键技术的拔尖创新人才，将采用项目制教学，打破传统学科界限。",
                "spider_from": "深职大官网",
                "url": "https://www.example.com/"
            },
            {
                "title": "金华职业技术大学入选国家级创新创业教育实践基地",
                "content": "在教育部最新公布的名单中，金华职业技术大学成功入选。学校表示将以此为契机，进一步完善创新创业课程体系，提升学生的双创能力。",
                "spider_from": "金职院新闻网",
                "url": ""
            }
        ]
    }

    print("正在渲染文档...")
    try:
        tpl.render(context)

        try:
            tpl.save(output_path)
            print(f"[成功] 文件已生成: {os.path.abspath(output_path)}")
        except PermissionError:
            timestamp = int(time.time())
            new_name = f"2026年第1期_生成测试_{timestamp}.docx"
            tpl.save(new_name)
            output_path = new_name
            print(f"[提示] 原文件名被占用，已另存为: {os.path.abspath(output_path)}")

        all_items = []
        for k in ("section_zhiku", "section_national", "section_provincial", "section_college"):
            all_items.extend(context.get(k) or [])

        postprocess_add_source_hyperlinks(output_path, all_items)
        print("[成功] 已完成后处理，来源已转为超链接")

    except Exception as e:
        print(f"[错误] 生成过程中发生异常: {e}")


if __name__ == "__main__":
    generate_weekly_report()
